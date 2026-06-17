"""
docx_parser.py — Layer 2b: extract text from Word (.docx) documents.

DOCX is actually easier than PDF — Word documents are structured XML, so
python-docx gives us the document hierarchy for free:
    - doc.paragraphs   → list of all paragraphs in order
    - para.style.name  → "Heading 1", "Heading 2", "Normal", "List Bullet" etc.
    - doc.tables       → list of all tables (rows × cells)

This means section detection is much more reliable for DOCX than PDF —
we can check paragraph.style.name instead of guessing from line length.
"""

import logging
import re
from io import BytesIO
from pathlib import Path
from typing import Union

from docx import Document
from docx.oxml.ns import qn

from .models import ParsedResume, FileType, ParseStatus

logger = logging.getLogger(__name__)


def parse(file_input: Union[str, Path, bytes, BytesIO]) -> ParsedResume:
    """
    Main entry point for DOCX parsing.
    Returns a ParsedResume with raw_text, sections, tables, and metadata.
    """
    raw_bytes = _to_bytes(file_input)
    filename  = _filename(file_input)
    warnings  = []

    try:
        doc = Document(BytesIO(raw_bytes))
    except Exception as e:
        logger.error(f"docx_parser: failed to open document — {e}")
        return ParsedResume(
            raw_text  = "",
            file_type = FileType.DOCX,
            status    = ParseStatus.FAILED,
            warnings  = [f"Could not open DOCX file: {e}"],
            filename  = filename,
        )

    # ── Extract paragraphs with style awareness ──────────────────────────────
    all_lines:      list[str]        = []
    sections:       dict[str, str]   = {}
    current_section = "header"
    section_buffer: list[str]        = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name if para.style else ""

        # Word's built-in heading styles: "Heading 1", "Heading 2", etc.
        # These are the most reliable section markers in a DOCX resume
        if _is_heading(style_name, text):
            # Save current section buffer
            if section_buffer:
                sections[current_section] = "\n".join(section_buffer).strip()
            current_section = text.lower().strip()
            section_buffer  = []
        else:
            section_buffer.append(text)
            all_lines.append(text)

    # Flush last section
    if section_buffer:
        sections[current_section] = "\n".join(section_buffer).strip()

    # ── Extract tables ───────────────────────────────────────────────────────
    table_texts, table_warnings = _extract_tables(doc)
    warnings.extend(table_warnings)

    # ── Build raw_text ───────────────────────────────────────────────────────
    raw_text = "\n".join(all_lines)
    if table_texts:
        raw_text += "\n\n[TABLE CONTENT]\n" + "\n".join(table_texts)
    raw_text = _clean_text(raw_text)

    # ── Status ───────────────────────────────────────────────────────────────
    if len(raw_text) >= 200:
        status = ParseStatus.SUCCESS
    elif len(raw_text) >= 50:
        status = ParseStatus.PARTIAL
        warnings.append(f"Low character count ({len(raw_text)}) — document may be empty or image-only")
    else:
        status = ParseStatus.FAILED
        warnings.append("Extracted text too short — document may contain only images")

    # DOCX doesn't have a meaningful "page count" — use paragraph count as proxy
    page_count_proxy = len(doc.paragraphs)

    return ParsedResume(
        raw_text   = raw_text,
        file_type  = FileType.DOCX,
        status     = status,
        sections   = sections,
        tables     = table_texts,
        warnings   = warnings,
        filename   = filename,
        page_count = page_count_proxy,
    )


def _is_heading(style_name: str, text: str) -> bool:
    """
    Determine if a paragraph is a section heading.

    Two ways a line can be a heading in a DOCX resume:
        1. It uses a Word heading style ("Heading 1", "Heading 2", "Title")
        2. It's short, ALL CAPS or Title Case, and matches our section keywords
           (many resumes use custom styles or no styles at all)
    """
    # Word's built-in heading styles
    if any(h in style_name for h in ["Heading", "Title", "Subtitle"]):
        return True

    # Fallback: short line that looks like a header
    if len(text) <= 40:
        # ALL CAPS headers: "EXPERIENCE", "SKILLS", "EDUCATION"
        if text.isupper() and len(text) > 2:
            return True
        # Title Case headers matching known section names
        normalised = re.sub(r"[^a-z\s]", "", text.lower()).strip()
        known = [
            "summary", "objective", "experience", "work experience",
            "education", "skills", "technical skills", "projects",
            "certifications", "awards", "languages", "references",
            "profile", "about me", "career history", "employment",
        ]
        if any(normalised == k or normalised.startswith(k) for k in known):
            return True

    return False


def _extract_tables(doc: Document) -> tuple[list[str], list[str]]:
    """
    Extract text from all tables in the document.

    Word tables are common in resumes for:
        - Two-column layouts (skills on left, contact info on right)
        - Skills matrices (competency × level grids)
        - Side-by-side experience blocks

    We extract all cell text and return it as individual strings.
    """
    warnings   = []
    cell_texts = []

    try:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        cell_texts.append(text)
    except Exception as e:
        warnings.append(f"Table extraction error: {e}")
        logger.warning(f"docx_parser._extract_tables: {e}")

    return cell_texts, warnings


def _clean_text(text: str) -> str:
    """Remove noise while preserving structure."""
    lines = text.splitlines()
    cleaned = []
    blank_count = 0

    for line in lines:
        line = line.strip()
        if not line:
            blank_count += 1
            if blank_count <= 2:
                cleaned.append("")
        else:
            blank_count = 0
            cleaned.append(line)

    return "\n".join(cleaned).strip()


def _to_bytes(file_input) -> bytes:
    if isinstance(file_input, (str, Path)):
        return Path(file_input).read_bytes()
    if isinstance(file_input, bytes):
        return file_input
    if isinstance(file_input, BytesIO):
        file_input.seek(0)
        return file_input.read()
    if hasattr(file_input, "read"):
        data = file_input.read()
        if hasattr(file_input, "seek"):
            file_input.seek(0)
        return data
    raise ValueError(f"Unsupported input type: {type(file_input)}")


def _filename(file_input) -> str:
    if isinstance(file_input, (str, Path)):
        return Path(file_input).name
    if hasattr(file_input, "name"):
        return Path(file_input.name).name
    return "unknown.docx"